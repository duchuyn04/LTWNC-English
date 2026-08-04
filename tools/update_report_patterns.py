from __future__ import annotations

import hashlib
import re
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\it\ltwnc")
INPUT = Path(r"C:\Users\juven\Downloads\NHOM 07_XAY DUNG WEBSITE HO TRO HOC TIENG ANH.docx")
OUTPUT = ROOT / "NHOM 07_XAY DUNG WEBSITE HO TRO HOC TIENG ANH_bo_sung_mau.docx"
EXPLAIN = ROOT / "explain"


PATTERNS = [
    ("command-report.md", "Command"),
    ("memento.md", "Memento"),
    ("factory.md", "Factory Method"),
    ("strategy.md", "Strategy"),
    ("observer.md", "Observer"),
    ("adapter.md", "Adapter"),
    ("decorator.md", "Decorator"),
    ("chain-of-responsibility.md", "Chain of Responsibility"),
]


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def strip_markdown(text: str) -> str:
    """Keep the report's inline-code convention while removing link syntax."""
    text = re.sub(r"\[([^\]]+)\]\([^\)]*\)", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("<br>", " ").replace("<br/>", " ")
    return re.sub(r"[ \t]+", " ", text).strip()


def parse_markdown(path: Path) -> list[tuple]:
    lines = path.read_text(encoding="utf-8").splitlines()
    events: list[tuple] = []
    paragraph: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            text = strip_markdown(" ".join(item.strip() for item in paragraph))
            if text:
                events.append(("paragraph", text))
            paragraph = []

    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            flush_paragraph()
            language = line[3:].strip().lower()
            index += 1
            code: list[str] = []
            while index < len(lines) and not lines[index].startswith("```"):
                code.append(lines[index].rstrip("\r"))
                index += 1
            if index < len(lines):
                index += 1
            events.append(("code", language, "\n".join(code)))
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            flush_paragraph()
            events.append(("heading", len(heading.group(1)), strip_markdown(heading.group(2))))
            index += 1
            continue

        if line.lstrip().startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            events.append(("table", table_lines))
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        if bullet:
            flush_paragraph()
            items: list[str] = []
            while index < len(lines):
                match = re.match(r"^\s*[-*]\s+(.*)$", lines[index])
                if not match:
                    break
                items.append(strip_markdown(match.group(1)))
                index += 1
            events.append(("list", items))
            continue

        numbered = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
        if numbered:
            flush_paragraph()
            items: list[str] = []
            while index < len(lines):
                match = re.match(r"^\s*\d+[.)]\s+(.*)$", lines[index])
                if not match:
                    break
                items.append(strip_markdown(match.group(1)))
                index += 1
            events.append(("list", items))
            continue

        if line.strip():
            paragraph.append(line.strip())
        else:
            flush_paragraph()
        index += 1

    flush_paragraph()
    return events


def normalize_heading(title: str) -> str:
    return re.sub(r"^(?:\d+\.)+\s*", "", title).strip()


def map_heading(title: str, pattern: str) -> str:
    clean = normalize_heading(title)
    lower = clean.lower()

    if "chỉ cần nhớ" in lower:
        return "Bối cảnh áp dụng"
    if "ví dụ dễ hiểu" in lower:
        return "Minh họa nhu cầu nghiệp vụ"
    if lower.startswith("code sau khi áp dụng"):
        return "Trình code triển khai thực tế trong project"
    if lower.startswith("trước khi") or lower.startswith("code trước khi"):
        return f"Cách triển khai trước khi áp dụng {pattern}"
    if lower.startswith("vì sao chọn") or lower.startswith("vì sao áp dụng") \
            or lower.startswith("vì sao dùng") or lower.startswith("vì sao đây") \
            or lower.startswith("biện luận việc lựa chọn"):
        return f"Biện luận nhu cầu áp dụng {pattern}"
    if "kết luận ngắn" in lower:
        return f"Kết quả đạt được sau khi áp dụng {pattern}"
    if lower.startswith("các vai trò") or lower.startswith("strategy nằm") \
            or lower.startswith("observer nằm") or lower.startswith("adapter nằm") \
            or lower.startswith("decorator nằm") or lower.startswith("chain nằm") \
            or lower.startswith("factory nằm"):
        return "Cấu trúc và vai trò trong project"
    if lower.startswith("thiết kế và triển khai"):
        return "Thiết kế và triển khai thực tế trong project"
    if lower.startswith("code sau"):
        return "Trình code triển khai thực tế trong project"
    if lower.startswith("ranh giới trách nhiệm"):
        return "Ranh giới trách nhiệm, giới hạn và kết quả"
    return clean


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_run_font(run, name: str = "Times New Roman", size: float = 13) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def add_body_paragraph(document: Document, text: str) -> Paragraph:
    paragraph = document.add_paragraph(style="time")
    run = paragraph.add_run(strip_markdown(text))
    set_run_font(run)
    return paragraph


def add_heading(document: Document, text: str, level: int = 3) -> Paragraph:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    return paragraph


def cell_text(cell: _Cell, text: str, *, code: bool = False, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    set_run_font(run, "Cascadia Mono" if code else "Times New Roman", 9.5 if code else 10.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code_table(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    lines = code.splitlines() or [""]
    cell_text(cell, lines[0], code=True)
    for line in lines[1:]:
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run(line)
        set_run_font(run, "Cascadia Mono", 9.5)
    document.add_paragraph(style="time")


def split_table_row(line: str) -> list[str]:
    content = line.strip().strip("|")
    return [strip_markdown(part.strip()) for part in content.split("|")]


def is_separator_row(row: Iterable[str]) -> bool:
    return all(re.fullmatch(r":?-{2,}:?", value.replace(" ", "")) for value in row)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines]
    rows = [row for row in rows if row and not is_separator_row(row)]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    usable_width = 6.35
    if column_count == 2:
        widths = [1.95, 4.4]
    elif column_count == 3:
        widths = [1.8, 2.5, 2.05]
    else:
        widths = [usable_width / column_count] * column_count
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for column_index, cell in enumerate(cells):
            cell.width = Inches(widths[column_index])
            value = row[column_index] if column_index < len(row) else ""
            cell_text(cell, value, bold=row_index == 0)
    document.add_paragraph(style="time")


def trim_body_after_command(document: Document) -> None:
    body = document._element.body
    start = None
    children = list(body)
    for index, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        paragraph = Paragraph(child, document)
        if paragraph.text.strip() == "Áp dụng mẫu thiết kế Command":
            start = index
            break
    if start is None:
        raise RuntimeError("Không tìm thấy vị trí bắt đầu mục Command trong tài liệu.")
    for child in children[start:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def enable_field_refresh(document: Document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def append_pattern(document: Document, path: Path, pattern: str) -> None:
    events = parse_markdown(path)
    document.add_paragraph(f"Áp dụng mẫu thiết kế {pattern}", style="Heading 2")

    for event in events:
        kind = event[0]
        if kind == "heading":
            level, title = event[1], event[2]
            if level == 1:
                continue
            mapped = map_heading(title, pattern)
            if not mapped:
                continue
            add_heading(document, mapped, 3)
        elif kind == "paragraph":
            add_body_paragraph(document, event[1])
        elif kind == "list":
            add_body_paragraph(document, " ".join(event[1]))
        elif kind == "code":
            language, code = event[1], event[2]
            if language == "mermaid":
                add_body_paragraph(document, "Sơ đồ quan hệ và luồng xử lý được giữ dưới dạng mã Mermaid để đối chiếu với thiết kế triển khai.")
            add_code_table(document, code)
        elif kind == "table":
            add_markdown_table(document, event[1])

    document.add_paragraph(style="time")


def main() -> None:
    if not INPUT.exists():
        raise FileNotFoundError(INPUT)
    source_hash = sha256(INPUT)
    shutil.copy2(INPUT, OUTPUT)
    document = Document(str(OUTPUT))
    trim_body_after_command(document)
    for filename, pattern in PATTERNS:
        append_pattern(document, EXPLAIN / filename, pattern)
    enable_field_refresh(document)
    document.save(str(OUTPUT))
    print(f"input={INPUT}")
    print(f"output={OUTPUT}")
    print(f"input_sha256={source_hash}")
    print(f"output_sha256={sha256(OUTPUT)}")
    print(f"output_size={OUTPUT.stat().st_size}")


if __name__ == "__main__":
    main()
